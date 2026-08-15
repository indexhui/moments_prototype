#!/usr/bin/env python3
"""Build the Cabinet Box Stack art-production brief as a polished DOCX."""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


WORKSPACE = Path("/Users/hugh/works/moment_prototype")
OUTPUT = WORKSPACE / "docs" / "art" / "CABINET_BOX_STACK_ART_BRIEF.docx"
SCREENSHOT_TUTORIAL = Path("/private/tmp/cabinet-box-tutorial-current.png")
SCREENSHOT_EMPTY = Path("/private/tmp/cabinet-box-gameplay-current.png")
SCREENSHOT_FAILURE = Path("/private/tmp/cabinet-box-failure-current.png")
SCREENSHOT_SEVEN = Path("/private/tmp/cabinet-box-seven-layer-current.png")

SKILL_DIR = Path(
    "/Users/hugh/.codex/plugins/cache/openai-primary-runtime/documents/"
    "26.805.11740/skills/documents"
)
sys.path.insert(0, str(SKILL_DIR / "scripts"))
from table_geometry import apply_table_geometry  # noqa: E402


INK = "243A3C"
ACCENT = "4F6B6B"
ACCENT_DARK = "344D51"
WARM = "B78146"
WARM_DARK = "755037"
BODY = "3E4748"
MUTED = "667274"
PALE = "EEF1ED"
PALE_WARM = "F7F1E7"
PALE_GOLD = "F4E8D3"
GRID = "CBD4D1"
WHITE = "FFFFFF"
RED = "9A4C43"
GREEN = "4F735F"

CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 90, "bottom": 90, "start": 120, "end": 120}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_border(cell, **edges) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name, edge_data in edges.items():
        tag = f"w:{edge_name}"
        edge = borders.find(qn(tag))
        if edge is None:
            edge = OxmlElement(tag)
            borders.append(edge)
        for key, value in edge_data.items():
            edge.set(qn(f"w:{key}"), str(value))


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_paragraph_left_border(paragraph, color: str, size: int = 18, space: int = 10) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = p_bdr.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        p_bdr.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), str(space))
    left.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def keep_row_together(row) -> None:
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    row.height = Twips(260)
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.keep_together = True


def set_font(run, *, size=None, bold=None, color=None, italic=None, font_name="Noto Sans TC") -> None:
    run.font.name = font_name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), font_name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), font_name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.italic = italic


def style_paragraph_runs(paragraph, *, size=11, color=BODY, bold=None) -> None:
    for run in paragraph.runs:
        set_font(run, size=size, color=color, bold=bold)


def add_text(doc, text: str = "", *, size=11, color=BODY, bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=6, line=1.25,
             keep_with_next=False, keep_together=False):
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    paragraph.paragraph_format.keep_with_next = keep_with_next
    paragraph.paragraph_format.keep_together = keep_together
    run = paragraph.add_run(text)
    set_font(run, size=size, color=color, bold=bold, italic=italic)
    return paragraph


def add_rich_paragraph(doc, parts, *, before=0, after=6, line=1.25,
                       align=WD_ALIGN_PARAGRAPH.LEFT, shading=None, left_border=None,
                       keep_together=False):
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    paragraph.paragraph_format.keep_together = keep_together
    paragraph.paragraph_format.left_indent = Pt(8) if left_border else Pt(0)
    paragraph.paragraph_format.right_indent = Pt(8) if shading else Pt(0)
    for text, options in parts:
        run = paragraph.add_run(text)
        set_font(
            run,
            size=options.get("size", 11),
            color=options.get("color", BODY),
            bold=options.get("bold", False),
            italic=options.get("italic", False),
            font_name=options.get("font_name", "Noto Sans TC"),
        )
    if shading:
        set_paragraph_shading(paragraph, shading)
    if left_border:
        set_paragraph_left_border(paragraph, left_border)
    return paragraph


def add_heading(doc, text: str, level: int = 1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    return paragraph


def add_bullet(doc, text: str, *, level=0, bold_lead: str | None = None):
    style_name = "List Bullet" if level == 0 else "List Bullet 2"
    paragraph = doc.add_paragraph(style=style_name)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.left_indent = Inches(0.375 if level == 0 else 0.625)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    if bold_lead and text.startswith(bold_lead):
        lead_run = paragraph.add_run(bold_lead)
        set_font(lead_run, size=11, bold=True, color=ACCENT_DARK)
        rest_run = paragraph.add_run(text[len(bold_lead):])
        set_font(rest_run, size=11, color=BODY)
    else:
        run = paragraph.add_run(text)
        set_font(run, size=11, color=BODY)
    return paragraph


def add_number(doc, text: str):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    run = paragraph.add_run(text)
    set_font(run, size=11, color=BODY)
    return paragraph


def format_table(table, *, header_fill=ACCENT, body_fill=WHITE, font_size=9.2,
                 header_font_size=9.2, alternate=True) -> None:
    table.style = "Table Grid"
    for row_index, row in enumerate(table.rows):
        keep_row_together(row)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            fill = header_fill if row_index == 0 else (
                PALE if alternate and row_index % 2 == 0 else body_fill
            )
            set_cell_shading(cell, fill)
            set_cell_border(
                cell,
                top={"val": "single", "sz": 5, "color": GRID},
                bottom={"val": "single", "sz": 5, "color": GRID},
                start={"val": "single", "sz": 5, "color": GRID},
                end={"val": "single", "sz": 5, "color": GRID},
            )
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.15
                paragraph.paragraph_format.keep_together = True
                for run in paragraph.runs:
                    set_font(
                        run,
                        size=header_font_size if row_index == 0 else font_size,
                        bold=row_index == 0,
                        color=WHITE if row_index == 0 else BODY,
                    )
    set_repeat_table_header(table.rows[0])


def add_table(doc, headers, rows, widths_dxa, *, font_size=9.2, header_fill=ACCENT,
              alternate=True):
    table = doc.add_table(rows=1, cols=len(headers))
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    apply_table_geometry(
        table,
        widths_dxa,
        table_width_dxa=CONTENT_DXA,
        indent_dxa=TABLE_INDENT_DXA,
        cell_margins_dxa=CELL_MARGINS,
    )
    format_table(
        table,
        header_fill=header_fill,
        font_size=font_size,
        alternate=alternate,
    )
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def set_crop(inline_shape, *, left=0, top=0, right=0, bottom=0) -> None:
    blip_fill = inline_shape._inline.graphic.graphicData.pic.blipFill
    src_rect = blip_fill.find(qn("a:srcRect"))
    if src_rect is None:
        src_rect = OxmlElement("a:srcRect")
        blip = blip_fill.find(qn("a:blip"))
        blip.addnext(src_rect)
    src_rect.set("l", str(int(left)))
    src_rect.set("t", str(int(top)))
    src_rect.set("r", str(int(right)))
    src_rect.set("b", str(int(bottom)))


def add_cropped_picture(run, path: Path, *, width_in: float, height_in: float,
                        crop_left=0, crop_top=0, crop_right=0, crop_bottom=0):
    shape = run.add_picture(str(path), width=Inches(width_in), height=Inches(height_in))
    set_crop(
        shape,
        left=crop_left,
        top=crop_top,
        right=crop_right,
        bottom=crop_bottom,
    )
    return shape


def add_figure_pair(doc, left_path: Path, right_path: Path, left_label: str, right_label: str):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.keep_together = True
    left_run = paragraph.add_run()
    add_cropped_picture(
        left_run,
        left_path,
        width_in=2.25,
        height_in=4.34,
        crop_left=3308,
        crop_top=6808,
        crop_right=3308,
        crop_bottom=10094,
    )
    gap = paragraph.add_run("   ")
    set_font(gap, size=4, color=WHITE)
    right_run = paragraph.add_run()
    add_cropped_picture(
        right_run,
        right_path,
        width_in=2.25,
        height_in=4.34,
        crop_left=3308,
        crop_top=6808,
        crop_right=3308,
        crop_bottom=10094,
    )
    caption = add_text(
        doc,
        f"{left_label}　　　　　　　　　　　　　{right_label}",
        size=8.6,
        color=MUTED,
        italic=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=0,
        after=8,
        line=1.0,
        keep_together=True,
    )
    return caption


def add_page_break(doc) -> None:
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def add_footer(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    label = paragraph.add_run("MOMENT · 疊箱子美術資產規格  |  ")
    set_font(label, size=8.5, color=MUTED)
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    set_font(run, size=8.5, color=MUTED)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_footer(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Noto Sans TC"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Noto Sans TC")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Noto Sans TC")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans TC")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BODY)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        1: (16, ACCENT_DARK, 18, 10),
        2: (13, ACCENT, 14, 7),
        3: (12, WARM_DARK, 10, 5),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Noto Sans TC"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Noto Sans TC")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Noto Sans TC")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans TC")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Bullet 2", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Noto Sans TC"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Noto Sans TC")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Noto Sans TC")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans TC")
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor.from_string(BODY)


def build_document() -> Document:
    for screenshot in (SCREENSHOT_TUTORIAL, SCREENSHOT_EMPTY, SCREENSHOT_FAILURE, SCREENSHOT_SEVEN):
        if not screenshot.exists():
            raise FileNotFoundError(f"Missing screenshot: {screenshot}")

    doc = Document()
    configure_styles(doc)
    doc.core_properties.title = "疊箱子小遊戲｜美術資產製作規格"
    doc.core_properties.subject = "Moment prototype cabinet box stack art production brief"
    doc.core_properties.author = "Moment Prototype Team"
    doc.core_properties.keywords = "Moment, art brief, cabinet, box stack, texture"

    # First-page memo masthead.
    add_text(doc, "ART PRODUCTION BRIEF · v0.1", size=10, color=WARM, bold=True, after=4, line=1.0)
    add_text(
        doc,
        "疊箱子小遊戲｜美術資產製作規格",
        size=27,
        color=INK,
        bold=True,
        after=4,
        line=1.0,
        keep_with_next=True,
    )
    add_text(
        doc,
        "Cabinet Box Stack — 將程式生成畫面改為美術繪製素材",
        size=13.5,
        color=MUTED,
        after=16,
        line=1.15,
        keep_with_next=True,
    )
    metadata = [
        ("對象", "2D 美術／UI 美術／技術美術"),
        ("平台", "Mobile Web，直式基準 393 × 852 px"),
        ("目前版本", "Three.js 箱塔 + Canvas 箱面標籤 + CSS 櫃體／教學圖"),
        ("文件日期", "2026-08-12"),
        ("建議做法", "美術提供可重複、可裁切的貼圖與前景圖層；程式保留互動與動態"),
    ]
    for label, value in metadata:
        add_rich_paragraph(
            doc,
            [
                (f"{label}：", {"size": 10.5, "bold": True, "color": ACCENT_DARK}),
                (value, {"size": 10.5, "color": BODY}),
            ],
            after=2,
            line=1.15,
            keep_together=True,
        )

    add_rich_paragraph(
        doc,
        [
            ("本文件的核心決策　", {"size": 11, "bold": True, "color": WARM_DARK}),
            (
                "不用請美術畫 14 個不同高度的完整箱塔。美術只需畫「櫃體、箱子三面材質、標籤／膠帶、教學示意、門片與必要圖示」；箱子的移動、切齊、裁掉、掉落、堆高、鏡頭與陰影由程式即時處理。",
                {"size": 11, "color": BODY},
            ),
        ],
        before=12,
        after=12,
        line=1.35,
        shading=PALE_GOLD,
        left_border=WARM,
        keep_together=True,
    )

    add_heading(doc, "一頁結論", 1)
    add_bullet(doc, "P0（本輪必做）：櫃體、櫃內三種表面、六組箱子皮膚、共用標籤紙、六色膠帶、教學示意圖、關門完成圖。", bold_lead="P0（本輪必做）：")
    add_bullet(doc, "P1（有餘裕再做）：方向圖示、重來圖示、星星、完美放置光圈、落空／掉落的小型特效。", bold_lead="P1（有餘裕再做）：")
    add_bullet(doc, "最重要的製作條件：箱體會被動態裁短，材質需可平鋪／裁切，不能把重要圖案貼在固定邊緣，也不要烘焙整體投影。", bold_lead="最重要的製作條件：")
    add_bullet(doc, "文字保留程式排版：箱名、代碼、進度、速度、回饋與按鈕文字不烘焙進圖，方便日後改字與多語。", bold_lead="文字保留程式排版：")

    add_heading(doc, "美術與程式的分工", 1)
    add_table(
        doc,
        ["美術負責", "程式負責"],
        [
            (
                "材質筆觸、色彩、磨損、紙箱摺線、膠帶、標籤紙、條碼裝飾、櫃體細節、門片、教學圖與圖示。",
                "箱體幾何、動態寬深、裁切碎片、堆疊位置、兩軸移動、速度、鏡頭、燈光、即時陰影、UI 文字與狀態切換。",
            ),
        ],
        [4680, 4680],
        font_size=9.6,
        header_fill=ACCENT_DARK,
        alternate=False,
    )

    # Current-state visual audit.
    add_heading(doc, "1. 現況畫面與玩法條件", 1)
    add_text(
        doc,
        "以下是 393 × 852 px 手機視窗的現況截圖。文件只把它當構圖與互動參考；目前的平塗、漸層與幾何感不是最終美術品質。",
        size=10.5,
        color=MUTED,
        after=8,
    )
    add_figure_pair(
        doc,
        SCREENSHOT_EMPTY,
        SCREENSHOT_SEVEN,
        "現況：初始移動中的箱子",
        "現況：7 層通關後的箱塔",
    )
    add_rich_paragraph(
        doc,
        [
            ("畫面基準　", {"size": 10.5, "bold": True, "color": ACCENT_DARK}),
            (
                "全螢幕 393 × 852；主櫃外框約 367 × 708；櫃內可視區約 343 × 682（皆為邏輯像素）。美術交付以 @3x 為主。",
                {"size": 10.5, "color": BODY},
            ),
        ],
        shading=PALE,
        left_border=ACCENT,
        after=8,
        line=1.25,
        keep_together=True,
    )

    add_heading(doc, "玩家會看到的狀態", 2)
    add_table(
        doc,
        ["狀態", "畫面行為", "美術需求", "備註"],
        [
            ("首次教學", "示意箱在層板上，顯示左右／深度兩方向", "教學示意圖 + 兩個方向圖示", "文字與按鈕由程式排"),
            ("移動中", "箱子沿 X 或深度軸往返", "同一箱體貼圖需適用兩軸", "每層速度 +18%"),
            ("成功放置", "重疊部分保留，超出部分切下", "貼圖需耐裁切；重要細節在安全區", "完美誤差約 ±6 px"),
            ("裁切／掉落", "超出碎片旋轉、下墜、淡出", "不另畫每種碎片；沿用箱體材質", "程式切幾何"),
            ("落空失敗", "整箱掉出，顯示失敗視窗", "可選落空小特效", "7 層前落空才失敗"),
            ("通關／結算", "7／10／14 層對應 1／2／3 星；門片關上", "關門圖、星星（P1）", "7 層後可提前完成"),
        ],
        [1440, 2340, 2880, 2700],
        font_size=8.6,
        header_fill=ACCENT,
    )

    # Art direction.
    add_heading(doc, "2. 美術方向與畫面分層", 1)
    add_heading(doc, "整體方向", 2)
    add_bullet(doc, "延續 Moment 的溫暖繪本／日常感：辦公櫃要有使用痕跡，但不要走寫實髒污或工業恐怖感。")
    add_bullet(doc, "箱子要像「公司資料箱」而不是搬家紙箱：有分類標籤、代碼、封箱膠帶、輕微壓痕與不同用途的色彩辨識。")
    add_bullet(doc, "輪廓與明暗必須在小螢幕仍清楚；前／側／頂三面要有穩定色差，讓玩家一眼看出箱體體積。")
    add_bullet(doc, "櫃內不宜過度花：背景細節的對比低於箱子，否則高速移動時會干擾對齊判斷。")

    add_heading(doc, "建議圖層順序（由後到前）", 2)
    layers = [
        ("01", "櫃內後壁／側壁／底面材質", "可平鋪貼圖；低對比"),
        ("02", "層板與底座", "接住第一箱；表面可受即時陰影"),
        ("03", "已放置箱／移動箱／掉落碎片", "共用六組箱皮；程式改尺寸與位置"),
        ("04", "接觸陰影、燈光、霧、鏡頭", "程式生成，不烘焙進素材"),
        ("05", "櫃框、底部擋板、名牌、門片", "前景 PNG 或獨立貼圖"),
        ("06", "方向、速度、進度、回饋字與按鈕", "UI 層；文字由程式"),
        ("07", "提示／教學／失敗／結算視窗", "UI 與插圖最上層"),
    ]
    add_table(
        doc,
        ["層", "內容", "製作方式"],
        layers,
        [960, 3600, 4800],
        font_size=9.1,
        header_fill=ACCENT_DARK,
    )

    add_heading(doc, "不可烘焙進圖的項目", 2)
    add_bullet(doc, "完整場景投影、接觸陰影、掉落模糊、移動殘影、成功光圈。")
    add_bullet(doc, "箱名、英文分類、代碼、層數、星等、速度、提示與按鈕文字。")
    add_bullet(doc, "固定 14 層的箱塔、固定裁切形狀、固定左右／深度位置。")

    add_rich_paragraph(
        doc,
        [
            ("安全區原則　", {"size": 10.5, "bold": True, "color": WARM_DARK}),
            (
                "紙箱的左右與深度可能逐層縮短。摺痕、髒點可散佈；膠帶與標籤需獨立成層。若必須畫在面材質上，重要內容至少離四邊 8%，且中心 60% 應能單獨成立。",
                {"size": 10.5, "color": BODY},
            ),
        ],
        shading=PALE_GOLD,
        left_border=WARM,
        before=8,
        after=10,
        line=1.3,
        keep_together=True,
    )

    # P0 asset list.
    add_heading(doc, "3. P0 必做資產清單", 1)
    add_text(
        doc,
        "建議以一個分層 PSD／AFPHOTO 為工作檔，再依下表匯出。若美術只交工作檔，也需依資產 ID 建立同名群組，讓程式／技術美術可批次切圖。",
        size=10.5,
        color=MUTED,
        after=8,
    )

    asset_rows = [
        ("ENV-01", "櫃體開啟前景框", "1", "1101 × 2124", "透明 PNG；中間挖空；含外框與底擋板"),
        ("ENV-02", "櫃內後壁材質", "1", "1024 × 1024", "方形、可平鋪、無烘焙陰影"),
        ("ENV-03", "櫃內側壁材質", "1", "1024 × 1024", "可與後壁同系列；色階稍深"),
        ("ENV-04", "櫃內底面／層板材質", "1", "1024 × 1024", "可平鋪；保留箱影可讀性"),
        ("ENV-05", "櫃門關閉正面", "1", "1029 × 2046", "透明／不透明皆可；中央 50% 可切成左右門片"),
        ("BOX-01~06", "六組箱子三面皮膚", "18", "見第 4 節", "每組 front／side／top；不含文字"),
        ("BOX-LABEL", "共用標籤紙底", "1", "1024 × 448", "透明 PNG；留白供程式排字／條碼"),
        ("BOX-TAPE-01~06", "六色膠帶", "6", "256 × 1024", "透明 PNG；長向可平鋪；避免亮斑固定在端點"),
        ("TUT-01", "教學對齊示意圖", "1", "936 × 462", "@3x；箱子 + 層板；不含說明文字"),
    ]
    add_table(
        doc,
        ["資產 ID", "內容", "數量", "交付尺寸 px", "重點"],
        asset_rows,
        [1500, 2100, 720, 1740, 3300],
        font_size=8.3,
        header_fill=ACCENT_DARK,
    )

    add_heading(doc, "櫃體製作說明", 2)
    add_number(doc, "先以 393 × 852 的現況畫面確認構圖；櫃框邏輯尺寸約 367 × 708，@3x 即 1101 × 2124。")
    add_number(doc, "ENV-01 中央必須透明，讓 3D 箱塔與櫃內表面在後方顯示；前景框不得遮到主要堆疊安全區。")
    add_number(doc, "ENV-05 需沿畫面正中央可無縫切成左右兩半，程式會讓兩門片由外向內關閉。中央接縫不可錯位。")
    add_number(doc, "櫃框／門片可以畫磨損、把手或名牌底；『ARCHIVE 03・文件櫃』文字仍由程式排。")

    add_heading(doc, "材質表面要求", 2)
    add_bullet(doc, "8-bit RGB／sRGB。一般材質不含透明；前景框、標籤紙與膠帶使用透明 PNG。")
    add_bullet(doc, "1024 方形貼圖四邊需真正無縫；避免中央唯一亮點、明顯方向性污漬或不可重複的刮痕。")
    add_bullet(doc, "不必提供法線貼圖；如美術熟悉 PBR，可額外提供同尺寸 grayscale roughness，程式端再評估。")

    # Boxes.
    add_heading(doc, "4. 六組資料箱皮膚規格", 1)
    add_rich_paragraph(
        doc,
        [
            ("共同結構　", {"size": 10.5, "bold": True, "color": ACCENT_DARK}),
            (
                "每組箱子只需三個面：front、side、top。底面不會成為視覺主角，可由程式使用深色紙板材質。標籤紙與膠帶獨立，不要畫死在三面底圖中。",
                {"size": 10.5, "color": BODY},
            ),
        ],
        shading=PALE,
        left_border=ACCENT,
        after=10,
        line=1.3,
        keep_together=True,
    )

    add_table(
        doc,
        ["面", "建議尺寸 px", "比例／方向", "內容限制"],
        [
            ("front", "1024 × 384", "橫向約 2.67:1", "主要紙板筆觸；中心預留標籤；左右可被裁短"),
            ("side", "512 × 384", "橫向約 1.33:1", "比正面深一階；深度方向可能被裁短"),
            ("top", "1024 × 512", "橫向 2:1", "摺線可見但低對比；膠帶另層；不可烘焙場景光"),
        ],
        [1440, 2160, 2040, 3720],
        font_size=9.0,
        header_fill=ACCENT,
    )

    add_heading(doc, "六組識別", 2)
    box_rows = [
        ("archive-a", "專案資料 A", "PROJECT FILE", "PRJ-A", "藍綠膠帶；中性暖棕紙箱"),
        ("archive-b", "專案資料 B", "PROJECT FILE", "PRJ-B", "草綠膠帶；與 A 保持同系列"),
        ("receipts", "收據備份", "ACCOUNTING", "ACC-24", "磚紅膠帶；紙箱可稍偏灰棕"),
        ("meeting", "會議附件", "MEETING DOCS", "MTG-07", "灰藍膠帶；乾淨、整齊"),
        ("samples", "樣品小物", "OFFICE SAMPLE", "SMP-12", "紫色膠帶；箱體可有輕微鼓起"),
        ("stationery", "備用文具", "STATIONERY", "ST-03", "橘棕膠帶；活潑但不要過亮"),
    ]
    add_table(
        doc,
        ["ID", "中文名稱", "分類字", "代碼", "視覺識別"],
        box_rows,
        [1500, 1860, 1800, 1200, 3000],
        font_size=8.7,
        header_fill=ACCENT_DARK,
    )

    add_heading(doc, "箱面標籤（共用底圖）", 2)
    add_bullet(doc, "BOX-LABEL 只畫紙張、邊框、折痕、黏貼痕與條碼／線段裝飾；上述中文、分類字、代碼由程式置入。")
    add_bullet(doc, "標籤底圖四周留 6% 透明或安全邊，縮小到手機畫面時仍要有清楚外框。")
    add_bullet(doc, "標籤寬高比固定，若箱子被裁到太窄，程式會縮小或隱藏標籤；美術不用另畫小尺寸版。")

    add_heading(doc, "箱體縮放／裁切測試", 2)
    add_bullet(doc, "front 必須在原寬 100%、75%、50% 三種裁切寬度都不露接縫。")
    add_bullet(doc, "side／top 必須在深度 100%、70%、40% 時仍成立；摺線與污漬不能只靠最外側才能看懂。")
    add_bullet(doc, "六組皮膚在高速移動時仍能辨認；色彩識別以膠帶為主，紙箱本體差異為輔。")

    # Tutorial, UI, VFX.
    add_heading(doc, "5. 教學、門片與 P1 圖示／特效", 1)
    add_heading(doc, "首次教學示意圖（TUT-01）", 2)
    add_text(
        doc,
        "現況教學圖由 CSS 多邊形拼出。美術改畫後，保留同樣的資訊：一個箱子浮在層板上方、可辨識正／側／頂三面、留出左右兩側放方向圖示。",
        size=10.5,
        color=BODY,
        after=6,
    )
    tutorial_figure = doc.add_paragraph()
    tutorial_figure.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tutorial_figure.paragraph_format.space_after = Pt(4)
    tutorial_figure.paragraph_format.keep_together = True
    tutorial_run = tutorial_figure.add_run()
    add_cropped_picture(
        tutorial_run,
        SCREENSHOT_TUTORIAL,
        width_in=3.85,
        height_in=2.52,
        crop_left=7600,
        crop_top=22400,
        crop_right=7600,
        crop_bottom=52000,
    )
    add_text(
        doc,
        "現況教學卡構圖參考（僅供資訊層級；不沿用現有平塗造型）",
        size=8.6,
        color=MUTED,
        italic=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=8,
        line=1.0,
    )
    add_bullet(doc, "輸出 936 × 462 px（@3x）；不含標題、說明文與按鈕。")
    add_bullet(doc, "背景可含低對比櫃內色塊；主要箱體需位於中央 60% 安全區。")
    add_bullet(doc, "方向箭頭建議另交，程式依當前移動軸顯示，不要把兩個箭頭烘焙進插圖。")

    add_heading(doc, "P1 圖示與小型 VFX", 2)
    add_table(
        doc,
        ["資產 ID", "內容", "尺寸 px", "格式／動畫方式"],
        [
            ("UI-RESET", "重新開始", "96 × 96", "透明 PNG 或 SVG；單色可變色優先"),
            ("UI-AXIS-X", "左右移動", "96 × 96", "透明 PNG 或 SVG"),
            ("UI-AXIS-Z", "深度斜向移動", "96 × 96", "透明 PNG 或 SVG"),
            ("UI-STAR-ON/OFF", "星星亮／暗", "96 × 96", "透明 PNG；兩態"),
            ("VFX-PERFECT", "完美放置光圈", "512 × 512", "透明 PNG；程式縮放＋淡出"),
            ("VFX-PLACE", "一般放置光圈", "512 × 512", "透明 PNG；色彩較弱"),
            ("VFX-DUST", "落空／裁切粉塵", "6 × 512 × 512", "橫向 spritesheet；選配"),
        ],
        [1800, 2520, 1560, 3480],
        font_size=8.7,
        header_fill=ACCENT,
    )

    # Delivery and acceptance.
    add_heading(doc, "6. 交付格式、命名與驗收", 1)
    add_rich_paragraph(
        doc,
        [
            ("開始畫之前需確認的唯一決策　", {"size": 10.5, "bold": True, "color": WARM_DARK}),
            (
                "櫃體最後採『手繪前景框 + 3D 內壁材質』（本文件建議）或改成整張 2D 櫃內插圖。若選後者，程式需另做透視與陰影適配；未確認前，先不要畫完整櫃內大圖。",
                {"size": 10.5, "color": BODY},
            ),
        ],
        shading=PALE_GOLD,
        left_border=WARM,
        before=2,
        after=8,
        line=1.3,
        keep_together=True,
    )
    add_heading(doc, "交付格式", 2)
    add_bullet(doc, "工作檔：PSD／PSB／AFPHOTO 皆可；圖層名稱需對應資產 ID，不合併文字參考層與正式底圖。")
    add_bullet(doc, "輸出檔：PNG-24（需要透明）或 JPG／PNG（不透明材質）；所有點陣圖以 sRGB 匯出。")
    add_bullet(doc, "材質貼圖建議 1024 的 2 次方尺寸；不可在匯出時自動裁透明邊，避免程式定位偏移。")
    add_bullet(doc, "每個檔案只放一個資產；不使用中文檔名、空格、版本日期或『final_final』。版本寫在資料夾或提交紀錄。")

    add_heading(doc, "命名範例", 2)
    tree_text = (
        "art/cabinet_box_stack/\n"
        "  env/cabinet_frame_open.png\n"
        "  env/cabinet_inner_back.png\n"
        "  env/cabinet_inner_side.png\n"
        "  env/cabinet_shelf.png\n"
        "  env/cabinet_closed_front.png\n"
        "  boxes/archive-a/front.png  side.png  top.png  tape.png\n"
        "  boxes/archive-b/front.png  side.png  top.png  tape.png\n"
        "  boxes/shared/label_panel.png\n"
        "  tutorial/tutorial_alignment.png\n"
        "  ui/axis_x.png  axis_z.png  reset.png\n"
        "  vfx/perfect_ring.png  place_ring.png"
    )
    code_paragraph = add_text(
        doc,
        tree_text,
        size=8.2,
        color=INK,
        after=6,
        line=1.0,
        keep_together=True,
    )
    set_paragraph_shading(code_paragraph, PALE)
    code_paragraph.paragraph_format.left_indent = Pt(10)
    code_paragraph.paragraph_format.right_indent = Pt(10)
    for run in code_paragraph.runs:
        set_font(run, size=8.2, color=INK, font_name="Menlo")

    add_heading(doc, "美術自檢清單", 2)
    checklist_rows = [
        ("尺寸", "@3x 與本文件尺寸一致；透明邊未被自動裁掉。"),
        ("可裁切", "箱面在 100%／75%／50% 寬度與 100%／70%／40% 深度皆無接縫。"),
        ("不烘焙", "沒有場景投影、完整光源、固定位置文字或固定箱塔。"),
        ("文字", "標籤紙留白；中文、英文分類與代碼不在正式圖檔內。"),
        ("辨識", "六組箱子以膠帶色可在手機尺寸快速辨認，色盲情境仍有明暗／紋理差。"),
        ("門片", "關門圖中央 50% 切分後可無縫閉合；左右邊緣不漏底。"),
        ("透明", "PNG 無白邊、黑邊、半透明髒邊與預乘 Alpha 問題。"),
        ("風格", "櫃內背景對比低於箱子；箱體三面明暗清楚但不寫實過度。"),
    ]
    add_table(
        doc,
        ["檢查項目", "通過條件"],
        checklist_rows,
        [2160, 7200],
        font_size=8.6,
        header_fill=ACCENT_DARK,
    )

    add_heading(doc, "程式整合驗收場景", 2)
    integration_checks = [
        "空櫃 + 第一箱移動：確認三面材質、方向圖示與標籤可讀。",
        "連續放置 4 層：確認六組皮膚循環、接觸陰影與高速移動辨識。",
        "故意偏放：確認被裁箱體與掉落碎片不拉伸、不露接縫。",
        "7 層通關：確認高塔鏡頭、前景框遮擋與完成按鈕。",
        "結算關門：確認左右門片中央接縫、星星與文字層級。",
    ]
    for check in integration_checks:
        paragraph = add_bullet(doc, check)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            set_font(run, size=8.6, color=BODY)

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
